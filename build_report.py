"""
Build Elera FYP1 Report - fresh document approach.
Copies the template, clears from Chapter 1 onward, and rewrites.
"""
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
import re, os, copy

ARTIFACTS = r"C:\Users\Mimdal\.gemini\antigravity\brain\31dec609-5d4d-46c0-ba41-f675980e1da3"
TEMPLATE = r"BIT FYP Report Template V3.1.docx"
OUTPUT = r"Elera_FYP1_Report_v4.docx"

INFO = {
    "title": "ELERA: AN INTELLIGENT WEB-BASED ARDUINO CIRCUIT DESIGN ASSISTANT",
    "name": "MOHAMAD IMAD ADDIN BIN JA'FAR",
    "id": "2127923",
    "supervisor": "DR AHMAD FATZILAH BIN MISMAN",
    "semester": "SEMESTER 2, 2025/2026",
    "date": "SEPTEMBER 2025",
    "course": "INFO 4401 REPORT",
}


def sf(run, name="Arial", size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def parse_md(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
    elements = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        if line.strip() == "---" or line.strip() == "":
            i += 1
            continue
        if line.startswith("#### "):
            elements.append(("h4", line[5:].strip().strip("*")))
            i += 1; continue
        if line.startswith("### "):
            elements.append(("h3", line[4:].strip().strip("*")))
            i += 1; continue
        if line.startswith("## "):
            elements.append(("h2", line[3:].strip().strip("*")))
            i += 1; continue
        if line.startswith("# "):
            elements.append(("h1", line[2:].strip().strip("*")))
            i += 1; continue
        if line.strip().startswith("[FIGURE") or line.strip().startswith("[TABLE"):
            elements.append(("figure", line.strip()))
            i += 1; continue
        if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
            tbl = []
            while i < len(lines) and "|" in lines[i]:
                s = lines[i].strip()
                if s and not all(c in "|-: " for c in s):
                    tbl.append([c.strip().strip("*") for c in s.split("|")[1:-1]])
                i += 1
            if tbl:
                elements.append(("table", tbl))
            continue
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            elements.append(("bullet", line.strip()[2:]))
            i += 1; continue
        m = re.match(r"^\s*(\d+)\.\s+(.+)", line)
        if m:
            elements.append(("numbered", m.group(1), m.group(2)))
            i += 1; continue
        m2 = re.match(r"^\s*\(([a-z])\)\s+(.+)", line)
        if m2:
            elements.append(("lettered", f"({m2.group(1)}) {m2.group(2)}"))
            i += 1; continue
        # Body paragraph
        para = line
        i += 1
        while i < len(lines):
            nl = lines[i].rstrip("\n")
            if (not nl.strip() or nl.startswith("#") or nl.startswith("- ") or
                nl.startswith("* ") or nl.strip().startswith("[FIGURE") or
                nl.strip().startswith("[TABLE") or re.match(r"^\s*\d+\.\s+", nl) or
                ("|" in nl and i+1 < len(lines) and "|" in lines[i+1])):
                break
            para += " " + nl.strip()
            i += 1
        if para.startswith("*(") and para.endswith(")*"):
            continue
        elements.append(("body", para.strip()))
    return elements


def render_bold(p, text, size=11):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            sf(r, size=size, bold=True)
        else:
            r = p.add_run(part)
            sf(r, size=size)


def write_elements(doc, elements):
    for item in elements:
        t = item[0]
        if t == "h1":
            continue
        elif t == "h2":
            txt = item[1]
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 2"]
            p.clear()
            r = p.add_run(txt.upper())
            sf(r, size=12, bold=True)
        elif t == "h3":
            txt = item[1]
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 3"]
            p.clear()
            r = p.add_run(txt)
            sf(r, size=12, bold=True)
        elif t == "h4":
            p = doc.add_paragraph()
            r = p.add_run(item[1])
            sf(r, size=11, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif t == "body":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            render_bold(p, item[1])
        elif t == "bullet":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.left_indent = Cm(1.27)
            pf.first_line_indent = Cm(-0.63)
            render_bold(p, f"\u2022  {item[1]}")
        elif t == "numbered":
            num = item[1]
            txt = item[2]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.left_indent = Cm(1.27)
            pf.first_line_indent = Cm(-0.63)
            render_bold(p, f"{num}.  {txt}")
        elif t == "lettered":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.left_indent = Cm(2.54)
            pf.first_line_indent = Cm(-0.63)
            render_bold(p, item[1])
        elif t == "figure":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(item[1])
            sf(r, size=11, italic=True)
        elif t == "table":
            rows = item[1]
            if len(rows) >= 2:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                # Add borders manually since template lacks Table Grid
                tbl_el = tbl._tbl
                tblPr = tbl_el.tblPr if tbl_el.tblPr is not None else docx.oxml.OxmlElement('w:tblPr')
                borders = docx.oxml.OxmlElement('w:tblBorders')
                for edge in ('top','left','bottom','right','insideH','insideV'):
                    el = docx.oxml.OxmlElement(f'w:{edge}')
                    el.set(qn('w:val'), 'single')
                    el.set(qn('w:sz'), '4')
                    el.set(qn('w:space'), '0')
                    el.set(qn('w:color'), '000000')
                    borders.append(el)
                tblPr.append(borders)
                for i, row in enumerate(rows):
                    for j, val in enumerate(row):
                        if j < len(tbl.columns):
                            cell = tbl.rows[i].cells[j]
                            cell.text = ""
                            r = cell.paragraphs[0].add_run(val)
                            sf(r, size=10, bold=(i == 0))
                doc.add_paragraph()


def main():
    print("Loading template...")
    doc = Document(TEMPLATE)

    # ── Fix cover page ──
    repls = [
        ("INFO 440X REPORT", INFO["course"]),
        ("THE AWARENESS OF PRIVACY ISSUES ON INTERNET", INFO["title"]),
        ("MOHD HANAS HUSIN", INFO["name"]),
        ("20014586", INFO["id"]),
        ("MOHD ABU BAKAR", INFO["supervisor"]),
    ]
    for p in doc.paragraphs[:30]:
        for old, new in repls:
            if old in p.text:
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
        for run in p.runs:
            run.text = run.text.replace("MARCH 2024", INFO["date"])
            run.text = run.text.replace("SEMESTER 2, 2023/2024", INFO["semester"])

    # ── Update Abstract ──
    with open(os.path.join(ARTIFACTS, "chapter1_introduction.md"), "r", encoding="utf-8") as f:
        content = f.read()
    m = re.search(r"# ABSTRACT\s*\n\n(.+?)(?:\n\n\*\(|\n\n---)", content, re.DOTALL)
    abstract_text = m.group(1).strip() if m else ""

    for i, p in enumerate(doc.paragraphs):
        if "Must be between 250" in p.text:
            p.clear()
            r = p.add_run(abstract_text)
            sf(r, size=11)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Clear the FYP1/FYP2 instruction paragraphs
            for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
                pp = doc.paragraphs[j]
                if "INFO 440" in pp.text or "FYP" in pp.text:
                    pp.clear()
            break

    # ── Remove everything from CHAPTER 1 title to end ──
    body = doc.element.body
    cut_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Title" and "CHAPTER 1" in p.text:
            cut_idx = i
            break

    if cut_idx is None:
        print("ERROR: CHAPTER 1 not found")
        return

    # Collect elements to remove
    to_remove = []
    for i in range(cut_idx, len(doc.paragraphs)):
        to_remove.append(doc.paragraphs[i]._element)

    # Also remove tables that come after cut point
    all_elements = list(body)
    first_cut_el = to_remove[0] if to_remove else None
    if first_cut_el is not None:
        cutting = False
        for el in all_elements:
            if el is first_cut_el:
                cutting = True
            if cutting:
                if el not in to_remove:
                    to_remove.append(el)

    for el in to_remove:
        try:
            body.remove(el)
        except ValueError:
            pass

    # ── Also remove FYP2-only preliminary pages ──
    # Find and remove Supervisor Approval, Certificate, Acknowledgement sections
    fyp2_titles = ["SUPERVISOR APPROVAL", "CERTIFICATE OF ORIGINALITY", "ACKNOWLEDGEMENT"]
    sections_to_remove = []
    in_section = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if any(t in txt for t in fyp2_titles):
            in_section = True
        if in_section:
            sections_to_remove.append(p._element)
            # Check if we hit the next major section (ABSTRACT, TABLE OF CONTENTS)
            if p.style and p.style.name == "Title" and txt in ["ABSTRACT", "TABLE OF CONTENTS"]:
                sections_to_remove.pop()  # keep the next title
                in_section = False

    for el in sections_to_remove:
        try:
            body.remove(el)
        except ValueError:
            pass

    # ── Write chapters ──
    print("Writing Chapter 1...")
    ch1 = parse_md(os.path.join(ARTIFACTS, "chapter1_introduction.md"))
    ch1 = [e for e in ch1 if not (e[0] == "h1" and ("ABSTRACT" in e[1].upper() or "CHAPTER" in e[1].upper() or "INTRODUCTION" in e[1].upper()))]
    # Skip the abstract paragraph (first body before first h2)
    filtered = []
    found_h2 = False
    for e in ch1:
        if e[0] == "h2":
            found_h2 = True
        if found_h2:
            filtered.append(e)
    # Add chapter title
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.clear()
    r = p.add_run("CHAPTER 1")
    sf(r, size=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    p2.style = doc.styles["Title"]
    p2.clear()
    r2 = p2.add_run("INTRODUCTION")
    sf(r2, size=14, bold=True)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    write_elements(doc, filtered)

    # Chapter 2
    print("Writing Chapter 2...")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.clear()
    r = p.add_run("CHAPTER 2")
    sf(r, size=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph()
    p2.style = doc.styles["Title"]
    p2.clear()
    r2 = p2.add_run("LITERATURE REVIEW")
    sf(r2, size=14, bold=True)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ch2 = parse_md(os.path.join(ARTIFACTS, "chapter2_literature_review.md"))
    ch2 = [e for e in ch2 if e[0] != "h1"]
    # Separate references section
    ch2_main = []
    ch2_refs = []
    in_refs = False
    for e in ch2:
        if e[0] == "h2" and "References" in e[1]:
            in_refs = True
            continue
        if in_refs:
            ch2_refs.append(e)
        else:
            ch2_main.append(e)
    write_elements(doc, ch2_main)

    # Chapter 3
    print("Writing Chapter 3...")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.clear()
    r = p.add_run("CHAPTER 3")
    sf(r, size=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph()
    p2.style = doc.styles["Title"]
    p2.clear()
    r2 = p2.add_run("METHODOLOGY")
    sf(r2, size=14, bold=True)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ch3 = parse_md(os.path.join(ARTIFACTS, "chapter3_methodology.md"))
    ch3 = [e for e in ch3 if e[0] != "h1"]
    write_elements(doc, ch3)

    # Chapter 4 (FYP2 placeholder - System Development)
    print("Writing Chapter 4 placeholder...")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.clear()
    r = p.add_run("CHAPTER 4")
    sf(r, size=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph()
    p2.style = doc.styles["Title"]
    p2.clear()
    r2 = p2.add_run("PROJECT DEVELOPMENT, IMPLEMENTATION AND EVALUATION")
    sf(r2, size=14, bold=True)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ch4 subheadings as placeholders
    ch4_headings = [
        ("h2", "4.0 INTRODUCTION"),
        ("body", "(To be completed in INFO 4402.)"),
        ("h2", "4.1 SYSTEM INTEGRATION"),
        ("body", "(To be completed in INFO 4402.)"),
        ("h2", "4.2 SYSTEM OUTPUT"),
        ("h3", "4.2.1 Administrator"),
        ("body", "(To be completed in INFO 4402.)"),
        ("h3", "4.2.2 User(s)"),
        ("body", "(To be completed in INFO 4402.)"),
        ("h2", "4.3 SYSTEM TESTING"),
        ("h3", "4.3.1 Test Plan"),
        ("body", "(To be completed in INFO 4402.)"),
        ("h3", "4.3.2 Enhancement"),
        ("body", "(To be completed in INFO 4402.)"),
    ]
    write_elements(doc, ch4_headings)

    # References
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.clear()
    r = p.add_run("REFERENCES")
    sf(r, size=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ref in ch2_refs:
        if ref[0] == "body":
            rp = doc.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = rp.paragraph_format
            pf.left_indent = Cm(1.27)
            pf.first_line_indent = Cm(-1.27)
            render_bold(rp, ref[1], size=11)

    # Appendices
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.clear()
    r = p.add_run("APPENDICES")
    sf(r, size=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ap = doc.add_paragraph()
    ar = ap.add_run("(Gantt chart and additional materials to be added)")
    sf(ar, size=11, italic=True)

    print(f"Saving {OUTPUT}...")
    doc.save(OUTPUT)
    print("Done! Report saved.")


if __name__ == "__main__":
    main()
